"""预处理模块 v2：PyMuPDF + 多模态LLM。

升级点：
1. PyMuPDF 替代 pypdf，提升文本提取质量
2. 多模态LLM 直接读图，替代 Tesseract OCR
3. 保留表格、标题等结构信息
"""
import base64
import logging
import re
from concurrent.futures import ThreadPoolExecutor, as_completed
from typing import Any

logger = logging.getLogger("tpaper.pipeline.preprocess_v2")


def preprocess_pdf_v2(content: bytes, use_vision: bool = True) -> dict:
    """PDF 预处理 v2：PyMuPDF 提取 + 多模态LLM识别扫描件。"""
    import fitz  # PyMuPDF
    
    doc = fitz.open(stream=content, filetype="pdf")
    pages = []
    
    for i, page in enumerate(doc):
        page_num = i + 1
        
        # 提取文本（保留布局）
        text = page.get_text("text")
        
        # 提取表格
        tables = []
        try:
            tab_finder = page.find_tables()
            for tab in tab_finder.tables:
                tables.append({
                    "rows": tab.extract(),
                    "bbox": tab.bbox
                })
        except Exception as e:
            logger.debug(f"第 {page_num} 页表格提取失败: {e}")
        
        # 提取图片数量
        images = page.get_images()
        
        # 判断是否需要多模态处理
        needs_multimodal = len(text.strip()) < 50 or len(images) > 0
        
        # 如果需要多模态，将页面渲染为图片
        image_b64 = None
        if needs_multimodal and use_vision:
            mat = fitz.Matrix(2, 2)  # 2倍放大
            pix = page.get_pixmap(matrix=mat)
            img_bytes = pix.tobytes("png")
            image_b64 = base64.b64encode(img_bytes).decode()
        
        pages.append({
            "page": page_num,
            "text": text,
            "tables": tables,
            "image_count": len(images),
            "needs_multimodal": needs_multimodal,
            "image_b64": image_b64,
            "mime": "image/png" if image_b64 else None,
        })
    
    doc.close()
    return {"pages": pages, "page_count": len(pages)}


def preprocess_docx_v2(content: bytes) -> dict:
    """DOCX 预处理 v2：python-docx 提取结构化内容。"""
    try:
        from docx import Document
        import io
        
        doc = Document(io.BytesIO(content))
        
        # 提取段落（保留样式）
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append({
                    "text": para.text,
                    "style": para.style.name if para.style else "Normal",
                    "is_heading": para.style.name.startswith("Heading") if para.style else False
                })
        
        # 提取表格
        tables = []
        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(cells)
            tables.append(rows)
        
        # 合并为结构化文本
        text_parts = []
        for p in paragraphs:
            prefix = "# " if p["is_heading"] else ""
            text_parts.append(f"{prefix}{p['text']}")
        
        for i, table in enumerate(tables):
            text_parts.append(f"\n[表格 {i + 1}]")
            for row in table:
                text_parts.append(" | ".join(row))
        
        text = "\n".join(text_parts)
        
        return {
            "pages": [{"page": 1, "text": text, "needs_multimodal": False}],
            "page_count": 1,
        }
    except ImportError:
        logger.warning("python-docx 未安装")
        return {"pages": [{"page": 1, "text": "", "needs_multimodal": True}], "page_count": 1}
    except Exception as e:
        logger.error(f"DOCX 解析失败: {e}")
        return {"pages": [{"page": 1, "text": "", "needs_multimodal": True}], "page_count": 1}


def preprocess_image_v2(content: bytes, mime: str) -> dict:
    """图片预处理 v2：直接进入多模态识别。"""
    b64 = base64.b64encode(content).decode()
    return {
        "pages": [{
            "page": 1,
            "text": "",
            "needs_multimodal": True,
            "image_b64": b64,
            "mime": mime,
            "tables": [],
            "image_count": 1,
        }],
        "page_count": 1,
    }


def split_by_content_boundaries(pages: list[dict]) -> list[dict]:
    """按内容边界分块，保留跨页题目的完整性。"""
    chunks = []
    current_chunk = {
        "pages": [],
        "text": "",
        "start_page": 1,
        "end_page": 1,
    }
    
    for page in pages:
        text = page["text"]
        
        # 检测内容边界（题目结束标记）
        if current_chunk["text"] and is_question_end(text):
            # 保存当前块
            current_chunk["end_page"] = page["page"] - 1
            chunks.append(current_chunk)
            current_chunk = {
                "pages": [],
                "text": "",
                "start_page": page["page"],
                "end_page": page["page"],
            }
        
        current_chunk["pages"].append(page)
        current_chunk["text"] += f"\n--- 第{page['page']}页 ---\n" + text
        current_chunk["end_page"] = page["page"]
    
    # 处理最后一个块
    if current_chunk["pages"]:
        chunks.append(current_chunk)
    
    return chunks


def is_question_end(text: str) -> bool:
    """检测题目结束标记"""
    patterns = [
        r'^\d+[.、．]\s*',  # 题号：1. 2. 3.
        r'^[一二三四五六七八九十]+[、．]\s*',  # 中文题号
        r'^第\d+题',  # 第1题
        r'^[（(]\d+[)）]',  # (1) (2)
        r'答案[：:]',  # 答案：
        r'解析[：:]',  # 解析：
    ]
    
    lines = text.strip().split('\n')
    if lines:
        first_line = lines[0].strip()
        return any(re.search(p, first_line) for p in patterns)
    
    return False


def split_with_overlap(pages: list[dict], window_size: int = 200) -> list[dict]:
    """带滑动窗口的分块，保留跨页上下文。"""
    chunks = split_by_content_boundaries(pages)
    
    # 添加前后重叠窗口
    enhanced_chunks = []
    all_text = "\n".join([p["text"] for p in pages])
    
    for i, chunk in enumerate(chunks):
        chunk_text = chunk["text"]
        
        # 添加前置窗口（前一块的末尾）
        if i > 0:
            prev_text = chunks[i-1]["text"]
            overlap_start = max(0, len(prev_text) - window_size)
            prefix = prev_text[overlap_start:]
            chunk_text = f"[上文]\n{prefix}\n\n[当前]\n{chunk_text}"
        
        # 添加后置窗口（后一块的开头）
        if i < len(chunks) - 1:
            next_text = chunks[i+1]["text"]
            overlap_end = min(len(next_text), window_size)
            suffix = next_text[:overlap_end]
            chunk_text = f"{chunk_text}\n\n[下文]\n{suffix}"
        
        enhanced_chunks.append({
            "text": chunk_text,
            "start_page": chunk["start_page"],
            "end_page": chunk["end_page"],
            "page_count": chunk["end_page"] - chunk["start_page"] + 1,
        })
    
    return enhanced_chunks


def preprocess_v2(source_file, use_vision: bool = True) -> dict:
    """预处理入口 v2。"""
    from app.config import settings
    from app.storage import get_storage
    
    storage = get_storage()
    content = storage.get(settings.source_files_namespace, source_file.storage_key)
    
    if source_file.mime_type == "application/pdf":
        return preprocess_pdf_v2(content, use_vision=use_vision)
    elif "wordprocessing" in source_file.mime_type:
        return preprocess_docx_v2(content)
    elif source_file.mime_type.startswith("image/"):
        return preprocess_image_v2(content, source_file.mime_type)
    else:
        return {"pages": [{"page": 1, "text": "", "needs_multimodal": True}], "page_count": 1}
