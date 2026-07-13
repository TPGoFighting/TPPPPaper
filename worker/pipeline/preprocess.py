"""预处理模块：PDF/DOCX/图片解析 + OCR。

对应 SPEC 14.1。
"""
import base64
import logging
from concurrent.futures import ThreadPoolExecutor, as_completed

logger = logging.getLogger("tpaper.pipeline.preprocess")


def preprocess_pdf(content: bytes, include_page_images: bool = False) -> dict:
    """PDF 预处理：提取文本，并为含图形的页面保留视觉上下文。

    仅以文字长度判断是否需要视觉模型会遗漏可提取文字页中的公式、
    表格或矢量图（例如算法题里的状态图）。启用视觉模型时，同时渲染
    含嵌入图像或矢量绘制的页面，让后续转写把这些信息与文本层合并。
    """
    extracted_pages = []
    try:
        import fitz  # pymupdf
        doc = fitz.open(stream=content, filetype="pdf")
        for i, page in enumerate(doc):
            text = page.get_text() or ""
            needs_multimodal = len(text.strip()) < 50
            page_info = {
                "page": i + 1,
                "text": text,
                "needs_multimodal": needs_multimodal,
            }

            # get_drawings 能识别 PDF 中的表格线、棋盘格、流程图等矢量
            # 元素；这些内容通常不在 get_text 的结果中。单个页面查询失败
            # 不应让整份文件无法处理。
            visual_count = 0
            if include_page_images:
                try:
                    visual_count = len(page.get_images(full=True)) + len(page.get_drawings())
                except Exception as exc:
                    logger.debug("第 %s 页视觉元素检测失败: %s", i + 1, exc)
                if needs_multimodal or visual_count:
                    pix = page.get_pixmap(matrix=fitz.Matrix(1.5, 1.5), alpha=False)
                    page_info["image_b64"] = base64.b64encode(pix.tobytes("png")).decode()
                    page_info["mime"] = "image/png"
                    page_info["has_visual_content"] = bool(visual_count)
            extracted_pages.append(page_info)
        doc.close()
    except ImportError:
        logger.warning("pymupdf 未安装，PDF 文本提取不可用")
        extracted_pages = [{"page": 1, "text": "", "needs_multimodal": True}]
    except Exception as e:
        logger.error(f"PDF 解析失败: {e}")
        extracted_pages = [{"page": 1, "text": "", "needs_multimodal": True, "error": str(e)}]

    # 对需要 OCR 的页面，并行 Tesseract OCR 提取文字
    needs_ocr = [p for p in extracted_pages if p.get("needs_multimodal") and not p.get("text", "").strip()]
    if needs_ocr:
        try:
            import fitz  # pymupdf
            import pytesseract
            from PIL import Image
            import io as _io

            doc = fitz.open(stream=content, filetype="pdf")
            mat = fitz.Matrix(2, 2)

            def _preprocess_image(img: "Image.Image") -> "Image.Image":
                """灰度化 + 二值化，提升 OCR 准确率。"""
                img = img.convert("L")  # 灰度
                # 自适应阈值二值化
                img = img.point(lambda p: 0 if p < 140 else 255)
                return img

            def _ocr_page(page_info: dict) -> dict:
                page_idx = page_info["page"] - 1
                if page_idx < 0 or page_idx >= len(doc):
                    return page_info
                page = doc[page_idx]
                pix = page.get_pixmap(matrix=mat)
                img_bytes = pix.tobytes("png")
                img = Image.open(_io.BytesIO(img_bytes))
                ocr_img = _preprocess_image(img)
                ocr_text = pytesseract.image_to_string(
                    ocr_img, lang="chi_sim+eng", config="--psm 6"
                )
                if ocr_text.strip():
                    page_info["text"] = ocr_text.strip()
                    page_info["needs_multimodal"] = False
                page_info["image_b64"] = base64.b64encode(img_bytes).decode()
                page_info["mime"] = "image/png"
                return page_info

            max_workers = min(8, len(needs_ocr))
            logger.info(f"并行 OCR: {len(needs_ocr)} pages, workers={max_workers}")
            with ThreadPoolExecutor(max_workers=max_workers) as pool:
                futures = {pool.submit(_ocr_page, p): p for p in needs_ocr}
                for future in as_completed(futures):
                    try:
                        result = future.result()
                        if result.get("text"):
                            logger.info(f"第 {result['page']} 页 OCR 提取: {len(result['text'])} chars")
                        else:
                            logger.warning(f"第 {result['page']} 页 OCR 未提取到文字")
                    except Exception as e:
                        p = futures[future]
                        logger.error(f"第 {p['page']} 页 OCR 失败: {e}")

            doc.close()
        except ImportError as e:
            logger.warning(f"OCR 依赖未安装: {e}")
        except Exception as e:
            logger.error(f"PDF OCR 失败: {e}")

    return {"pages": extracted_pages, "page_count": len(extracted_pages)}


def preprocess_docx(content: bytes) -> dict:
    """DOCX 预处理：提取段落、表格和媒体。"""
    try:
        from docx import Document
        import io
        doc = Document(io.BytesIO(content))
        lines = []
        for p in doc.paragraphs:
            if p.text.strip():
                lines.append(p.text.strip())
        for ti, table in enumerate(doc.tables):
            lines.append(f"\n[表格 {ti + 1}]")
            for ri, row in enumerate(table.rows):
                cells = [cell.text.strip() for cell in row.cells]
                lines.append(" | ".join(cells))
        text = "\n".join(lines)
        return {
            "pages": [{"page": 1, "text": text, "needs_multimodal": False}],
            "page_count": 1,
        }
    except ImportError:
        logger.warning("python-docx 未安装")
        return {"pages": [{"page": 1, "text": "", "needs_multimodal": True}], "page_count": 1}
    except Exception as e:
        logger.error(f"DOCX 解析失败: {e}")
        return {"pages": [{"page": 1, "text": "", "needs_multimodal": True}], "page_count": 1}


def preprocess_image(content: bytes, mime: str) -> dict:
    """图片预处理：直接进入多模态识别。"""
    b64 = base64.b64encode(content).decode()
    return {
        "pages": [{"page": 1, "text": "", "needs_multimodal": True, "image_b64": b64, "mime": mime}],
        "page_count": 1,
    }


def preprocess(source_file, include_page_images: bool = False) -> dict:
    """根据类型分发预处理。"""
    from app.config import settings
    from app.storage import get_storage

    storage = get_storage()
    content = storage.get(settings.source_files_namespace, source_file.storage_key)

    if source_file.mime_type == "application/pdf":
        return preprocess_pdf(content, include_page_images=include_page_images)
    elif "wordprocessing" in source_file.mime_type:
        return preprocess_docx(content)
    elif source_file.mime_type.startswith("image/"):
        return preprocess_image(content, source_file.mime_type)
    else:
        return {"pages": [{"page": 1, "text": "", "needs_multimodal": True}], "page_count": 1}
