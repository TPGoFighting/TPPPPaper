"""TPaper 进程内任务处理（本地开发模式，无需 Redis/Worker）。

对应 SPEC 第 14 节 AI 处理流水线：
预处理 → 两阶段模型处理 → 网页生成 → 净化 → 校验恢复

从 worker/main.py 提取处理逻辑，改为使用相对导入，
并通过 process_queued_papers() 轮询数据库处理排队论文。
"""
import asyncio
import base64
import json
import logging
from datetime import datetime, timezone

from .config import settings
from .database import SessionLocal
from .models import (
    ModelProfile, Paper, PaperDraft, ProcessingJob, SourceFile,
)
from .security import decrypt_secret, sanitize_html, sanitize_css
from .storage import get_storage
from .adapters import (
    OpenAICompatibleAdapter, build_document_prompt, build_extraction_prompt,
    build_presentation_prompt,
)

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(name)s: %(message)s",
)
logger = logging.getLogger("tpaper.processing")


# ── 文档预处理（SPEC 14.1）──

def preprocess_pdf(content: bytes) -> dict:
    """PDF 预处理：提取文本、版面块和图片。

    MVP 简化实现：尝试用 pypdf 提取文本；文本不足时标记需要多模态。
    """
    extracted_pages = []
    try:
        from pypdf import PdfReader
        import io
        reader = PdfReader(io.BytesIO(content))
        for i, page in enumerate(reader.pages):
            text = page.extract_text() or ""
            extracted_pages.append({
                "page": i + 1,
                "text": text,
                "needs_multimodal": len(text.strip()) < 50,  # 文本不足
            })
    except ImportError:
        logger.warning("pypdf 未安装，PDF 文本提取不可用")
        extracted_pages = [{"page": 1, "text": "", "needs_multimodal": True}]
    except Exception as e:
        logger.error(f"PDF 解析失败: {e}")
        extracted_pages = [{"page": 1, "text": "", "needs_multimodal": True, "error": str(e)}]
    return {"pages": extracted_pages, "page_count": len(extracted_pages)}


def preprocess_docx(content: bytes) -> dict:
    """DOCX 预处理：提取段落、表格和媒体。"""
    try:
        from docx import Document
        import io
        doc = Document(io.BytesIO(content))
        lines = []
        # 提取段落
        for p in doc.paragraphs:
            if p.text.strip():
                lines.append(p.text.strip())
        # 提取表格
        for ti, table in enumerate(doc.tables):
            lines.append(f"\n[表格 {ti + 1}]")
            for ri, row in enumerate(table.rows):
                cells = [cell.text.strip() for cell in row.cells]
                lines.append(" | ".join(cells))
        text = "\n".join(lines)
        return {
            "pages": [{"page": 1, "text": text, "needs_multimodal": False}],
            "page_count": 1,
        }
    except ImportError:
        logger.warning("python-docx 未安装")
        return {"pages": [{"page": 1, "text": "", "needs_multimodal": True}], "page_count": 1}
    except Exception as e:
        logger.error(f"DOCX 解析失败: {e}")
        return {"pages": [{"page": 1, "text": "", "needs_multimodal": True}], "page_count": 1}


def preprocess_image(content: bytes, mime: str) -> dict:
    """图片预处理：直接进入多模态识别。"""
    b64 = base64.b64encode(content).decode()
    return {
        "pages": [{"page": 1, "text": "", "needs_multimodal": True, "image_b64": b64, "mime": mime}],
        "page_count": 1,
    }


def preprocess(source_file: SourceFile) -> dict:
    """根据类型分发预处理。"""
    storage = get_storage()
    content = storage.get(settings.source_files_namespace, source_file.storage_key)

    if source_file.mime_type == "application/pdf":
        return preprocess_pdf(content)
    elif "wordprocessing" in source_file.mime_type:
        return preprocess_docx(content)
    elif source_file.mime_type.startswith("image/"):
        return preprocess_image(content, source_file.mime_type)
    else:
        return {"pages": [{"page": 1, "text": "", "needs_multimodal": True}], "page_count": 1}


def _flatten_preprocessed_text(preprocessed: dict) -> str:
    chunks: list[str] = []
    for page in preprocessed.get("pages", []):
        text = (page.get("text") or "").strip()
        if text:
            chunks.append(text)
    return "\n\n".join(chunks).strip()


def build_fallback_document(title: str, preprocessed: dict, mode: str) -> dict:
    """在模型不可用时生成一版可审核草稿，保证上传-审核-发布流程不断线。"""
    source_text = _flatten_preprocessed_text(preprocessed)
    lines = [line.strip() for line in source_text.splitlines() if line.strip()]
    if not lines:
        lines = ["暂未从源文件中提取到足够文字，请在审核页补充题干与答案。"]

    question_lines = lines[:12]
    questions = []
    for index, line in enumerate(question_lines, start=1):
        is_blank = "____" in line or "（" in line and "）" in line
        question_type = "fill_blank" if is_blank else "subjective"
        question = {
            "id": f"q{index}",
            "number": index,
            "type": question_type,
            "stem": line,
            "score": 0,
            "source_page": 1,
            "confidence": 0.45,
            "needs_review": True,
            "is_ai_generated": mode == "lecture_to_quiz",
            "explanation": "这是系统在模型不可用时生成的兜底草稿，请人工审核。",
        }
        if question_type == "fill_blank":
            question["acceptable_answers"] = [["请补充答案"]]
            question["match_rule"] = "contains"
        else:
            question["reference_answer"] = "请在审核页补充参考答案。"
            question["scoring_points"] = ["人工确认题意", "补全答案或评分要点"]
        questions.append(question)

    return {
        "title": title or "未命名试卷",
        "language": "zh-CN",
        "metadata": {
            "generated_by": "local_fallback",
            "mode": mode,
            "review_required": True,
            "source_excerpt": source_text[:1000],
        },
        "sections": [
            {
                "id": "s1",
                "title": "待审核题目",
                "source_page": 1,
                "question_ids": [q["id"] for q in questions],
            }
        ],
        "questions": questions,
    }


def ensure_publishable_document(document: dict) -> dict:
    """补齐模型漏填的题型必需字段，让草稿可审核、可发布。"""
    questions = document.get("questions") or []
    for question in questions:
        qtype = question.get("type")
        question.setdefault("needs_review", False)
        if qtype in ("single_choice", "multi_choice"):
            options = question.get("options") or []
            if not options:
                question["options"] = [
                    {"key": "A", "text": "请补充选项 A"},
                    {"key": "B", "text": "请补充选项 B"},
                ]
                question["correct_keys"] = ["A"]
                question["needs_review"] = True
            elif not question.get("correct_keys"):
                first_key = options[0].get("key", "A") if isinstance(options[0], dict) else "A"
                question["correct_keys"] = [first_key]
                question["needs_review"] = True
        elif qtype == "true_false" and question.get("true_false_answer") is None:
            question["true_false_answer"] = True
            question["needs_review"] = True
        elif qtype == "fill_blank" and not question.get("acceptable_answers"):
            question["acceptable_answers"] = [["请补充答案"]]
            question["match_rule"] = "contains"
            question["needs_review"] = True
        elif qtype == "subjective":
            has_answer = (
                question.get("reference_answer")
                or question.get("scoring_points")
                or question.get("explanation")
            )
            if not has_answer:
                question["reference_answer"] = "请在审核页补充参考答案。"
                question["scoring_points"] = ["人工确认题意", "补全答案或评分要点"]
                question["needs_review"] = True
    return document


# ── 两阶段模型处理（SPEC 14.2）──

async def stage1_extract(adapter: OpenAICompatibleAdapter, preprocessed: dict, vision_adapter: OpenAICompatibleAdapter | None = None) -> list[dict]:
    """第一阶段：来源提取。多模态失败时自动降级为纯文本。"""
    extracted = []
    for page in preprocessed["pages"]:
        if page.get("needs_multimodal") and page.get("image_b64") and vision_adapter:
            # 多模态提取
            prompt = f"提取第 {page['page']} 页的所有文字、题目结构和图表信息。返回 JSON。"
            result = await vision_adapter.chat_with_image(
                prompt, page["image_b64"], page.get("mime", "image/png")
            )
            if not result.success:
                # 降级：多模态失败时用纯文本适配器
                logger.warning(f"多模态提取失败，降级为纯文本: {result.error}")
                msgs = build_extraction_prompt(page.get("text", ""), page["page"])
                result = await adapter.chat(msgs, response_format_json=True)
        else:
            msgs = build_extraction_prompt(page.get("text", ""), page["page"])
            result = await adapter.chat(msgs, response_format_json=True)

        if result.success:
            try:
                extracted.append(json.loads(result.content))
            except json.JSONDecodeError:
                extracted.append({"page": page["page"], "text": result.content, "confidence": 0.5})
        else:
            logger.warning(f"第 {page['page']} 页提取失败: {result.error}")
            extracted.append({"page": page["page"], "text": "", "error": result.error})
    return extracted


async def stage2_generate_document(
    adapter: OpenAICompatibleAdapter, extracted: list[dict], mode: str
) -> dict:
    """第二阶段：生成 PaperDocument。"""
    msgs = build_document_prompt(extracted, mode)
    result = await adapter.chat(msgs, response_format_json=True)
    if not result.success:
        raise RuntimeError(f"文档生成失败: {result.error}")
    try:
        return json.loads(result.content)
    except json.JSONDecodeError as e:
        # SPEC 14.4: JSON 不合法时自动进行一次结构修复
        logger.warning(f"JSON 解析失败，尝试修复: {e}")
        # 简单修复：提取第一个 { 到最后一个 }
        content = result.content
        start = content.find("{")
        end = content.rfind("}")
        if start >= 0 and end > start:
            return json.loads(content[start : end + 1])
        raise


# ── 网页生成（SPEC 14.3）──

async def stage3_generate_presentation(adapter: OpenAICompatibleAdapter, document: dict) -> tuple[str, str]:
    """使用模板渲染器生成受控 HTML 与 CSS。"""
    from .presentation import render_paper
    return render_paper(document)


# ── 主处理流程 ──

async def process_paper(paper_id: int, source_file_id: int) -> None:
    """处理一份资料的完整流水线。"""
    db = SessionLocal()
    try:
        paper = db.get(Paper, paper_id)
        source = db.get(SourceFile, source_file_id)
        if not paper or not source:
            logger.error(f"Paper {paper_id} 或 SourceFile {source_file_id} 不存在")
            return

        # 获取活跃模型 Profile。没有配置时使用本地兜底生成器，前端仍可审核发布。
        profile = db.query(ModelProfile).filter(ModelProfile.is_active.is_(True)).first()
        text_adapter = None
        vision_adapter = None
        if profile:
            api_key = decrypt_secret(profile.encrypted_api_key) if profile.encrypted_api_key else ""
            text_adapter = OpenAICompatibleAdapter(
                base_url=profile.base_url,
                api_key=api_key,
                model=profile.text_model,
                timeout=profile.timeout_seconds,
                allow_private_network=profile.allow_private_network,
            )
            vision_adapter = OpenAICompatibleAdapter(
                base_url=profile.base_url,
                api_key=api_key,
                model=profile.multimodal_model,
                timeout=profile.timeout_seconds,
                allow_private_network=profile.allow_private_network,
            )

        # 创建任务记录
        job = ProcessingJob(
            paper_id=paper_id,
            model_profile_id=profile.id if profile else None,
            job_type="parse",
            status="running",
            stage="preprocessing",
            idempotency_key=f"paper-{paper_id}-parse",
        )
        db.add(job)
        paper.status = "parsing"
        db.commit()

        # 阶段 1：预处理
        logger.info(f"[Paper {paper_id}] 预处理...")
        preprocessed = preprocess(source)
        job.total_pages = preprocessed["page_count"]
        job.stage = "extracting"
        source.page_count = preprocessed["page_count"]
        db.commit()

        if text_adapter:
            try:
                # 阶段 2：来源提取
                logger.info(f"[Paper {paper_id}] 来源提取...")
                extracted = await stage1_extract(
                    text_adapter, preprocessed,
                    vision_adapter if profile and profile.supports_vision else None,
                )
                job.stage = "generating_document"
                job.current_page = preprocessed["page_count"]
                db.commit()

                # 阶段 3：生成 PaperDocument
                logger.info(f"[Paper {paper_id}] 生成结构化文档...")
                paper.status = "modeling"
                db.commit()
                document = await stage2_generate_document(text_adapter, extracted, paper.mode)
            except Exception as model_error:
                logger.warning(f"[Paper {paper_id}] 模型生成失败，改用本地兜底: {model_error}")
                document = build_fallback_document(paper.title, preprocessed, paper.mode)
                job.error_code = type(model_error).__name__
                job.error_message = str(model_error)[:500]
        else:
            logger.warning(f"[Paper {paper_id}] 未配置模型 Profile，使用本地兜底")
            job.stage = "generating_document"
            job.current_page = preprocessed["page_count"]
            db.commit()
            document = build_fallback_document(paper.title, preprocessed, paper.mode)

        # 阶段 4：网页生成
        job.stage = "generating_presentation"
        db.commit()
        logger.info(f"[Paper {paper_id}] 生成网页...")
        document = ensure_publishable_document(document)
        html, css = await stage3_generate_presentation(text_adapter, document)

        # 阶段 5：净化
        job.stage = "sanitizing"
        db.commit()
        clean_html, _ = sanitize_html(html)
        clean_css, _ = sanitize_css(css, scope_selector="")
        validation_errors: list[str] = []
        try:
            from .schemas import PaperDocument

            validation_errors = PaperDocument.model_validate(document).semantic_validate()
        except Exception as e:
            validation_errors = [f"文档结构错误: {e}"]
        is_valid = len(validation_errors) == 0

        last_draft = (
            db.query(PaperDraft)
            .filter(PaperDraft.paper_id == paper_id)
            .order_by(PaperDraft.version.desc())
            .first()
        )
        next_draft_version = (last_draft.version + 1) if last_draft else 1

        # 创建草稿
        draft = PaperDraft(
            paper_id=paper_id,
            version=next_draft_version,
            document=document,
            presentation_html=clean_html,
            theme_css=clean_css,
            is_valid=is_valid,
            validation_result={"errors": validation_errors, "is_valid": is_valid},
        )
        db.add(draft)
        db.flush()

        paper.current_draft_id = draft.id
        paper.status = "pending_review"
        job.status = "succeeded"
        job.stage = "done"
        job.call_summary = {
            "model": profile.text_model if profile else "local_fallback",
            "pages_processed": preprocessed["page_count"],
        }
        db.commit()
        logger.info(f"[Paper {paper_id}] 处理完成，进入待审核")

    except Exception as e:
        logger.exception(f"[Paper {paper_id}] 处理失败")
        paper = db.get(Paper, paper_id)
        if paper:
            paper.status = "failed"
        if 'job' in locals():
            job.status = "failed"
            job.error_code = type(e).__name__
            job.error_message = str(e)[:500]
        db.commit()
    finally:
        db.close()


# ── 轮询处理（本地开发模式，无需 Redis）──

async def process_queued_papers() -> None:
    """轮询数据库处理排队的论文（本地开发模式，无需 Redis）。"""
    logger.info("启动进程内任务处理器，轮询排队论文...")
    while True:
        db = SessionLocal()
        try:
            queued = db.query(Paper).filter(Paper.status == "queued").all()
            for paper in queued:
                if paper.source_file_id:
                    await process_paper(paper.id, paper.source_file_id)
        except Exception as e:
            logger.error(f"轮询处理异常: {e}")
        finally:
            db.close()
        await asyncio.sleep(3)
